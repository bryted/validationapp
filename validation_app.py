import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
import io

# --- Streamlit Config ---
st.set_page_config(page_title="Data Quality Validation App", layout="wide")
st.title("📊 Data Quality Validation App")

# --- Step 1: Upload Files ---
st.header("Step 1: Upload Key and Data Files")
key_file = st.file_uploader("🔑 Upload the KEY Excel file", type=["xlsx"], key="key")
data_file = st.file_uploader("📂 Upload the DATA Excel file", type=["xlsx"], key="data")

# --- Globals ---
row_only_fields = []
row_or_column_fields = ["VisitType", "ReNoSchool", "RChildType", "RHhType", "EducStatus"]

answer_map = {}
key_description, data_sheets = None, None

# --- Step 2: Process Files ---
if key_file and data_file:
    with st.spinner("Reading and processing files..."):
        try:
            key_description = pd.read_excel(key_file, sheet_name=None)
            data_sheets = pd.read_excel(data_file, sheet_name=None)
        except Exception as e:
            st.error(f"❌ Error reading files: {e}")
            st.stop()

        # --- Process answer_list ---
        answer_list_df = key_description.get("answer_list", pd.DataFrame())
        if not answer_list_df.empty:
            answer_list_df.dropna(how="all", inplace=True)
            answer_list_df.dropna(axis=1, how="all", inplace=True)
            grouped = answer_list_df.groupby("Field Name Eng for partner")
            for field, group in grouped:
                if pd.isna(field):
                    continue
                all_values = []
                for _, row in group.iterrows():
                    if field in row_only_fields:
                        values = row.drop(labels=["Language", "Field Name Eng for partner"]).dropna().tolist()
                    elif field in row_or_column_fields:
                        row_vals = row.drop(labels=["Language", "Field Name Eng for partner"]).dropna().tolist()
                        col_vals = group.columns[2:][~row[2:].isna()].tolist()
                        values = row_vals + col_vals
                    else:
                        values = group.columns[2:][~row[2:].isna()].tolist()
                    all_values.extend([str(v).strip() for v in values if str(v).strip()])
                unique_values = list(dict.fromkeys(all_values))
                if unique_values:
                    answer_map[field] = unique_values

    # --- Step 3: Country & Language ---
    st.header("Step 2: Select Country and Language")
    country_choice = st.selectbox("🌍 Select Country", options=["GHA", "CIV"], index=0)
    language_choice = st.selectbox("🗣️ Select Language", options=["EN", "FR"], index=0)

    # --- Step 4: Run Validation ---
    if st.button("🚦 Run Validation"):
        with st.spinner("Running validation checks..."):
            # --- Load description sheet robustly ---
            if isinstance(key_description, dict):
                normalized = {s.strip().lower(): s for s in key_description.keys()}
                if "description" in normalized:
                    description_df = key_description[normalized["description"]]
                else:
                    st.error(f"❌ Could not find 'description' sheet. Available sheets: {list(key_description.keys())}")
                    st.stop()
            else:
                description_df = key_description

            # Normalize column names
            description_df.columns = description_df.columns.str.strip().str.replace("\n", " ", regex=True)

            if "For Mars KPI reporting" not in description_df.columns:
                st.error("❌ Column 'For Mars KPI reporting' not found in description sheet")
                st.write("Available columns:", description_df.columns.tolist())
                st.stop()

            required_meta = description_df[
                description_df["For Mars KPI reporting"].isin(["Y", country_choice])
            ]
            expected_fields = required_meta["Field Name Eng for partner"].dropna().unique().tolist()
            field_types = dict(
                zip(description_df["Field Name Eng for partner"], description_df["Type of variable in the table"])
            )

            data_issues = []

            def log_issue(sheet, field, row, issue_type, description):
                data_issues.append(
                    {"Sheet": sheet, "Field": field, "Row": row, "Issue Type": issue_type, "Description": description}
                )

            def validate_sheet(df, sheet_name):
                msgs = {
                    "empty_sheet": {
                        "EN": "Sheet is present but contains no data",
                        "FR": "La feuille est présente mais ne contient aucune donnée",
                    },
                    "completely_empty": {
                        "EN": "The variable '{}' is completely empty.",
                        "FR": "La variable '{}' est complètement vide.",
                    },
                    "missing_required": {
                        "EN": "Field is required but missing",
                        "FR": "Champ requis mais manquant",
                    },
                    "expected_date": {
                        "EN": "Expected format is DD-MM-YYYY",
                        "FR": "Format attendu : JJ-MM-AAAA",
                    },
                    "expected_numeric": {
                        "EN": "Expected a numeric value",
                        "FR": "Une valeur numérique était attendue",
                    },
                    "invalid_value": {
                        "EN": "'{}' not in allowed values",
                        "FR": "'{}' ne fait pas partie des valeurs autorisées",
                    },
                    "duplicate_combo": {
                        "EN": "Duplicate entries found based on combination of: {}",
                        "FR": "Doublons détectés selon la combinaison : {}",
                    },
                    "cond_re_no": {
                        "EN": "ChldAvble is 0, ReNoAvble must have a value",
                        "FR": "ChldAvble est 0, ReNoAvble doit contenir une valeur",
                    },
                    "cond_hh_re_no": {
                        "EN": "HH_Avble is 0, HH_ReNoAvble must have a value",
                        "FR": "HH_Avble est 0, HH_ReNoAvble doit contenir une valeur",
                    },
                }

                if df.empty:
                    log_issue(sheet_name, None, None, "Empty Sheet", msgs["empty_sheet"][language_choice])
                    return

                for field in expected_fields:
                    if field not in df.columns:
                        continue
                    col_data = df[field].replace(r"^\s*$", np.nan, regex=True)

                    # --- Conditional rules ---
                    if sheet_name == "D" and field == "ReNoAvble" and "ChldAvble" in df.columns:
                        for idx, val in col_data.items():
                            if str(df.loc[idx, "ChldAvble"]).strip() == "0" and pd.isna(val):
                                log_issue(sheet_name, field, idx, "Conditional Rule", msgs["cond_re_no"][language_choice])
                    elif sheet_name == "B-C" and field == "HH_ReNoAvble" and "HH_Avble" in df.columns:
                        for idx, val in col_data.items():
                            if str(df.loc[idx, "HH_Avble"]).strip() == "0" and pd.isna(val):
                                log_issue(
                                    sheet_name, field, idx, "Conditional Rule", msgs["cond_hh_re_no"][language_choice]
                                )

                    # --- Missing entire column ---
                    if col_data.isnull().all():
                        log_issue(
                            sheet_name,
                            field,
                            None,
                            "Missing Value",
                            msgs["completely_empty"][language_choice].format(field),
                        )
                        continue

                    # --- Row-by-row checks ---
                    for idx, val in col_data.items():
                        if pd.isna(val):
                            log_issue(sheet_name, field, idx, "Missing Value", msgs["missing_required"][language_choice])
                            continue

                        if field in field_types:
                            if "date" in str(field_types[field]).lower():
                                try:
                                    pd.to_datetime(val, format="%d-%m-%Y")
                                except Exception:
                                    log_issue(
                                        sheet_name, field, idx, "Date Format Error", msgs["expected_date"][language_choice]
                                    )
                            elif "float" in str(field_types[field]).lower():
                                try:
                                    pd.to_numeric(val)
                                except Exception:
                                    log_issue(
                                        sheet_name, field, idx, "Type Error", msgs["expected_numeric"][language_choice]
                                    )

                        if field in answer_map:
                            if str(val).strip().lower() not in [
                                str(v).strip().lower() for v in answer_map[field]
                            ]:
                                log_issue(
                                    sheet_name,
                                    field,
                                    idx,
                                    "Invalid Value",
                                    msgs["invalid_value"][language_choice].format(val),
                                )

                # --- Duplicate check ---
                keys = [f for f in ["ChldID", "FarmerID", "VisitType", "EndDateActivity"] if f in df.columns]
                if len(keys) >= 2:
                    composite_key = df[keys].fillna("").astype(str).agg("|".join, axis=1)
                    dupes = composite_key[composite_key.duplicated()]
                    if not dupes.empty:
                        log_issue(
                            sheet_name,
                            keys[0],
                            None,
                            "Duplicate",
                            msgs["duplicate_combo"][language_choice].format(", ".join(keys)),
                        )

            # Run validation for selected sheets
            for sheet_name in ["P", "B-C", "D", "E_Com", "E_Hho", "E_Chd"]:
                if sheet_name in data_sheets:
                    validate_sheet(data_sheets[sheet_name], sheet_name)

        # --- Step 5: Show Results ---
        st.subheader("Validation Summary")
        if data_issues:
            summary_df = pd.DataFrame(data_issues)

            def aggregate_duplicates(group):
                descriptions = group["Description"].tolist()
                if group.name[2] == "Duplicate":
                    values = ", ".join(
                        sorted(set(v.split(":")[-1].strip() for v in descriptions if ":" in v))
                    )
                    return pd.Series(
                        {"Issue Count": len(group), "Example Description": f"Duplicate IDs found: {values}"}
                    )
                return pd.Series({"Issue Count": len(group), "Example Description": descriptions[0]})

            grouped_summary = (
                summary_df.groupby(["Sheet", "Field", "Issue Type"], as_index=False)
                .apply(aggregate_duplicates)
            )

            # --- Display in expanders ---
            for sheet in grouped_summary["Sheet"].unique():
                with st.expander(f"📄 Issues in Sheet: {sheet}", expanded=False):
                    st.dataframe(grouped_summary[grouped_summary["Sheet"] == sheet])

            # --- Banner to remind downloads ---
            st.info("📥 Scroll down to download the validation report in Word or Excel format.")

            # --- Word download ---
            word_buffer = io.BytesIO()
            word_doc = Document()
            word_doc.add_heading("MARS Data Quality Report", 0)

            # Summary section
            for sheet in grouped_summary["Sheet"].unique():
                word_doc.add_heading(f"Sheet: {sheet} - Summary", level=1)
                issues = grouped_summary[grouped_summary["Sheet"] == sheet]
                table = word_doc.add_table(rows=1, cols=4)
                table.style = "Light List Accent 1"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Field"
                hdr_cells[1].text = "Issue Type"
                hdr_cells[2].text = "Issue Count"
                hdr_cells[3].text = "Example Description"
                for _, row in issues.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row["Field"])
                    row_cells[1].text = str(row["Issue Type"])
                    row_cells[2].text = str(row["Issue Count"])
                    row_cells[3].text = str(row["Example Description"])

            # Detailed section
            for sheet in summary_df["Sheet"].unique():
                word_doc.add_heading(f"Sheet: {sheet} - Detailed Issues", level=2)
                issues = summary_df[summary_df["Sheet"] == sheet]
                table = word_doc.add_table(rows=1, cols=5)
                table.style = "Light Grid Accent 1"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Row"
                hdr_cells[1].text = "Field"
                hdr_cells[2].text = "Issue Type"
                hdr_cells[3].text = "Description"
                hdr_cells[4].text = "Sheet"
                for _, row in issues.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row["Row"])
                    row_cells[1].text = str(row["Field"])
                    row_cells[2].text = str(row["Issue Type"])
                    row_cells[3].text = str(row["Description"])
                    row_cells[4].text = str(row["Sheet"])

            word_doc.save(word_buffer)
            word_buffer.seek(0)

            st.download_button(
                label="📥 Download Word Report",
                data=word_buffer,
                file_name=f"Validation_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            # --- Excel download ---
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                summary_df.to_excel(writer, sheet_name="All Issues", index=False)
                grouped_summary.to_excel(writer, sheet_name="Grouped Summary", index=False)
            excel_buffer.seek(0)

            st.download_button(
                label="📊 Download Excel Report",
                data=excel_buffer,
                file_name=f"Validation_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        else:
            st.success("✅ No validation issues found!")
